"""
parsers/pdf_parser.py

PDF/DOCX attachment parser based on:
  1) table extractor   (pdfplumber, docx2pdf)
  2) text-block extractor for "RELAÇÃO DE ITENS" PDFs
  3) table-to-item parser (heuristics for licitação "Item / Quantidade / Unidade / Descrição" tables)

It produces ItemExtraido items compatible with the merger/aggregator pipeline.

Key entrypoint:
  parse_attachment(file_path, doc_type, project_root, tables_out_dir, parsed_out_dir, debug)

--- Relação de Itens text format ---
These PDFs are NOT table-based. Each item is a free-text block with the pattern:

    <N> - <Name>
    Descrição Detalhada: <long text …>
    …
    Quantidade Total: <int>   Quantidade Mínima Cotada: <int>
    …
    Unidade de Fornecimento: <unit>
    …

Detection heuristic: if the page text contains the phrase "RELAÇÃO DE ITENS" or
"Relação de Itens" we skip the table path entirely and use the block parser.
"""
from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Any


def _safe_relpath(path: Path, root: Path) -> str:
    """Return ``path`` relative to ``root``, falling back to ``str(path)``."""
    try:
        return str(path.resolve().relative_to(root.resolve()))
    except Exception:
        return str(path)


from models import ItemExtraido, ResultadoLicitacao
from extractor import process_file, write_tables_json


# =============================================================================
# Shared helpers
# =============================================================================

def _norm(x: Any) -> str:
    """Normalise any value to a single-space-separated stripped string."""
    if x is None:
        return ""
    return re.sub(r"\s+", " ", str(x)).strip()


def _flatten_table(table: list[list[Any]]) -> list[list[str]]:
    """Convert a raw pdfplumber table (which may contain None cells) to
    a list of normalised string rows."""
    return [[_norm(c) for c in row] for row in table or []]


def _table_text(rows: list[list[str]]) -> str:
    """Concatenate table rows to a single text block for regex-based metadata
    extraction. Non-empty cells are joined with `` | `` within each row."""
    lines = []
    for row in rows:
        if any(row):
            lines.append(" | ".join([c for c in row if c]))
    return "\n".join(lines)


def _parse_int_ptbr(s: str) -> Optional[int]:
    """Parse a Brazilian-formatted integer string, returning None on failure.

    Handles:
      - Thousand-separator dots:   "1.178.040" → 1178040
      - Decimal comma (truncated): "1.5,3"     → 15  (integer part only)
      - Plain integers:            "280"        → 280

    Returns None rather than raising so callers can distinguish "not a number"
    from "zero".
    """
    s = _norm(s)
    if not s:
        return None
    # Remove thousand-separator dots and spaces.
    s = s.replace(".", "").replace(" ", "")
    # If there is a comma, treat everything before it as the integer part.
    if re.match(r"^\d+,\d+$", s):
        s = s.split(",")[0]
    if re.match(r"^\d+$", s):
        try:
            return int(s)
        except ValueError:
            return None
    return None



# ── Lote parsing ──────────────────────────────────────────────────────────────

# Generic "LOTE" matcher used in both table and text-block parsers.
# We keep it permissive for the token, and format/normalise downstream.
_RE_LOTE_GENERIC = re.compile(r"\bLOTE\s*(?:N[º°]\s*)?([A-Z]?\d{1,3}|ÚNICO|UNICO)\b", re.I)


def _format_lote(token: str, suffix: str = "") -> str:
    """Normalise a lote token to a stable label.

    Examples:
      - "1"      -> "LOTE 01"
      - "01"     -> "LOTE 01"
      - "ÚNICO"  -> "LOTE ÚNICO"
      - "UNICO"  -> "LOTE ÚNICO"
      - "A1"     -> "LOTE A01"

    If ``suffix`` is provided (e.g. "PRINCIPAL", "RESERVADO ME/EPP"), it is
    appended in uppercase.
    """
    tok = _norm(token).upper()
    if tok in ("ÚNICO", "UNICO"):
        core = "ÚNICO"
    else:
        # Optional letter prefix (rare, but some editais use "A1", "B2"...)
        m = re.fullmatch(r"([A-Z])(\d{1,3})", tok)
        if m:
            core = m.group(1) + m.group(2).zfill(2)
        else:
            # Digits only
            core = re.sub(r"\D", "", tok)
            core = core.zfill(2) if core else tok

    suf = _norm(suffix)
    if suf:
        suf = re.sub(r"^[\-–:]+", "", suf).strip()
    if suf:
        return f"LOTE {core} {suf.upper()}"
    return f"LOTE {core}"


def _extract_lote_from_row(row: list[str]) -> Optional[str]:
    """Extract a lote label from a *single* table row, if present.

    Important: lote headers sometimes appear *mid-table* (e.g. a table contains
    the tail of a previous lot, then a "LOTE 09 ..." row, then the header for
    the next lot). Therefore, lote detection must be row-based and sequential.

    Returns:
        Normalised lote label (e.g. "LOTE 09 PRINCIPAL") or None.
    """
    if not row or not any(row):
        return None
    txt = _norm(" ".join([c for c in row if c]))
    if not txt:
        return None

    # Fast check
    if "LOTE" not in txt.upper():
        return None

    # Prefer matches that start at the beginning of the row (common for lote titles).
    m = re.search(r"^\s*LOTE\s*(?:N[º°]\s*)?([A-Z]?\d{1,3}|ÚNICO|UNICO)\b(.*)$", txt, re.I)
    if m:
        return _format_lote(m.group(1), m.group(2))

    # Fallback: allow lote mention later in the row only if it's very near the start.
    m2 = _RE_LOTE_GENERIC.search(txt)
    if m2 and m2.start() <= 3:
        return _format_lote(m2.group(1), txt[m2.end():])

    return None


def _row_is_header(row: list[str]) -> Optional[dict[str, int]]:
    """Return a header_map if the row looks like a table header, else None."""
    hm = _detect_header_indices(row)
    if "item" in hm and "obj" in hm:
        return hm
    return None

# ── Item number parsing ────────────────────────────────────────────────────────

# Matches hierarchical item numbers with 2–4 levels, e.g. "1.2", "13.2.1", "5.1.1.1".
# Major part: up to 4 digits; each subsequent segment: 1–2 digits.
# This is intentionally strict to avoid misclassifying quantities like "1.178.040" as item IDs.
_RE_HIER_ITEM = re.compile(r"^(\d{1,4})(\.\d{1,2}){1,3}$")

# Matches flat item numbers like "1", "007", "12345".
_RE_FLAT_ITEM = re.compile(r"^\d{1,5}$")


def _parse_item_id(s: str) -> Optional[tuple[int, Optional[str]]]:
    """Parse a cell value as an item identifier.

    Returns:
        ``(major, sub)`` for hierarchical items:
            "1.2"     → (1, "2")
            "13.2.1"  → (13, "2.1")
            "5.1.1.1" → (5, "1.1.1")
        ``(n, None)``  for flat items ("7" → (7, None)).
        ``None``       if the string is not a recognisable item number.

    Also handles PDF line-split artefacts like "5.1.1. 1" (space/newline before
    the last digit after a dot), which are normalised to "5.1.1.1" before matching.
    """
    s = _norm(s)
    # Some PDFs render item numbers with leading/trailing punctuation, e.g. ".4", "4.", "• 12".
    s = re.sub(r"^[\s\.\-–•]+", "", s)
    s = re.sub(r"[\.\-–]+$", "", s)
    # Fix PDF line-break artefact: "5.1.1. 1" → "5.1.1.1"
    s = re.sub(r"\.\s+(\d)", r".\1", s)

    m = _RE_HIER_ITEM.match(s)
    if m:
        parts = s.split(".")
        return int(parts[0]), ".".join(parts[1:])
    if _RE_FLAT_ITEM.match(s):
        return int(s), None
    return None



def _is_item_number(s: str) -> bool:
    """Return True when the string looks like an item identifier."""
    return _parse_item_id(s) is not None


def _is_section_title(text: str) -> bool:
    """Return True when the text looks like a section/group header (not an item).

    The original heuristic (uppercase ratio) was too aggressive for licitação tables
    like this PDF, where legitimate continuation lines can be ALL CAPS
    (e.g. "NA LOCAÇÃO COM FORNECIMENTO DE MATERIAL").

    We keep the uppercase check, but require the text to also look like a *real*
    section label via keywords and short length.
    """
    if not text:
        return False

    t = _norm(text)
    if len(t) > 60:
        return False
    if ":" in t:
        return False

    # Only treat as a "section" if it contains a typical section keyword.
    if not re.search(r"\b(GRUPO|SUBGRUPO|SEÇÃO|SECAO|CATEGORIA|LOTE|ANEXO|TOTAL)\b", t.upper()):
        return False

    alpha = [c for c in t if c.isalpha()]
    if not alpha:
        return False

    return sum(c.isupper() for c in alpha) / len(alpha) > 0.65


def _looks_like_code(s: str) -> bool:
    """Return True when the string appears to be a numeric/code identifier.

    A "code" here means something like a CATMAT/CATSERV code or a price code —
    a string composed entirely of digits, dots, hyphens and slashes, at least
    4 characters long. Such strings are penalised when scoring object candidates.
    """
    s = _norm(s)
    return bool(re.fullmatch(r"[\d\.\-\/]+", s)) and len(s) >= 4


def _compact_header_token(s: str) -> str:
    """Compact a header cell to a token for robust matching.

    This PDF often splits header words across lines/cells (e.g. "ITE
M",
    "LOC
AL"). By removing whitespace/punctuation we can still match them as
    "ITEM" and "LOCAL".
    """
    return re.sub(r"[^A-Z0-9]", "", _norm(s).upper())


def _detect_header_indices(header_row: list[str]) -> dict[str, int]:
    """Map column roles to their indices by inspecting a candidate header row.

    Same intent as before, but more tolerant to split headers like "ITE M".

    Returns keys among: item, qtd, und, obj.
    """
    idx: dict[str, int] = {}
    for i, cell in enumerate(header_row):
        if not cell:
            continue

        raw = _norm(cell).upper().strip()
        tok = _compact_header_token(cell)

        # ITEM / Nº ITEM variants (tok catches "ITE M" -> "ITEM")
        if tok in ("ITEM", "ITEM.", "NITEM", "NOITEM", "NºITEM", "N°ITEM", "N°", "Nº") or raw.startswith("ITEM"):
            idx.setdefault("item", i)

        # Quantidade
        # Quantidade ("QTD", "QUANT.", "QUANTIDADE")
        if (
            "QTD" in tok
            or "QUANTIDADE" in raw
            or tok.startswith("QUANT")
            or tok in ("QT", "QT.")
        ):
            idx.setdefault("qtd", i)

        # Unidade
        if any(k in tok for k in ("UND", "UNID", "UNIDADE")) or tok in ("UN", "UN."):
            idx.setdefault("und", i)

        # Descrição / Objeto
        if any(k in raw for k in ("DESCRI", "ESPECIF", "OBJETO", "SERVIÇO", "SERVICO", "MATERIAL", "PRODUTO")):
            idx.setdefault("obj", i)

    return idx


def _pick_obj(row: list[str], preferred_idx: int) -> str:
    """Select the best object/description text from a data row.

    Scoring rules (higher = better candidate):
      +len(text)  — prefer longer strings
      +20         — bonus for the mapped ``obj`` column
      -10         — penalty for strings that look like codes

    Cells that are item numbers, pure integers, or contain "R$" are excluded
    from consideration entirely.

    Falls back to the raw value at ``preferred_idx`` when no alphabetic
    candidate is found.
    """
    candidates: list[tuple[int, str]] = []
    for j, c in enumerate(row):
        if not c:
            continue
        if _is_item_number(c):
            continue
        if _parse_int_ptbr(c) is not None:
            continue
        if "R$" in c:
            continue
        if re.search(r"[A-Za-zÀ-ú]", c):
            score = len(c)
            if j == preferred_idx:
                score += 20
            if _looks_like_code(c):
                score -= 10
            # Avoid choosing the unit-of-supply cell as the object when the
            # header map is off by one (common on continuation pages).
            # Examples seen in PDFs: objeto="unid" / objeto="und".
            tok2 = _compact_header_token(c)
            if tok2 in {
                "UN", "UNID", "UND", "UNIDADE", "UNIDADES",
                "CX", "CAIXA", "PCT", "PACOTE", "KIT",
                "SERV", "SERVICO", "SERVIÇO",
                "M", "M2", "M3", "L", "KG",
            }:
                score -= 50
            candidates.append((score, c))
    if candidates:
        return max(candidates, key=lambda x: x[0])[1]
    if preferred_idx < len(row):
        return row[preferred_idx]
    return ""


def _extract_qty_with_header(row: list[str], qtd_idx: int, item_num: int) -> int:
    """Extract the quantity value from a data row.

    First tries the mapped ``qtd`` column; if that fails or returns 0,
    searches within ±4 columns around ``qtd_idx`` for any positive integer
    that is not equal to ``item_num`` (to avoid confusing the item number
    itself with the quantity when they happen to be in adjacent columns).

    Args:
        row:      Full data row.
        qtd_idx:  Column index mapped to ``"qtd"`` by the header detector.
        item_num: The item's own number, used to exclude false positives.

    Returns:
        Extracted quantity as a positive int, or 0 if none found.
    """
    if 0 <= qtd_idx < len(row):
        cell = row[qtd_idx]
        # Guard: when the header map is wrong, the quantity window may hit the
        # item-id cell (e.g. "5.1.1. 1"), which would parse as 5111.
        if not ("." in _norm(cell) and _is_item_number(cell)):
            q = _parse_int_ptbr(cell)
            if q is not None and q > 0:
                return q
    best_q: Optional[int] = None
    for j in range(max(0, qtd_idx - 4), min(len(row), qtd_idx + 5)):
        if j == qtd_idx:
            continue
        cell = row[j]
        if "." in _norm(cell) and _is_item_number(cell):
            continue
        q = _parse_int_ptbr(cell)
        if q is not None and q > 0:
            if q != item_num:
                return q
            elif best_q is None:
                best_q = q
    return best_q if best_q is not None else 0


def _extract_unit_with_header(row: list[str], und_idx: int, obj_text: str = "") -> str:
    """Extract the unit-of-supply text from a data row.

    Strategy:
      1) Try the mapped ``und`` column (when available).
      2) Otherwise, scan a small window around ``und_idx``.
      3) If ``und_idx`` is unknown (-1) or no good candidate is found, scan the whole row.

    Unlike the older "first match wins" fallback, we score candidates and pick the
    best one. This avoids classic failures where the *object* value (e.g. "Container")
    is mistakenly returned as the unit when the header map is missing/shifted.

    Important nuance for licitação PDFs:
    units are often written with package qualifiers that include digits,
    e.g. "CX. C/\n50" ("caixa com 50"). Those must be accepted; otherwise the
    unit becomes empty and the item may be dropped by sanitisation.

    Normalises common PDF quirks like split dots: "SERV\n." -> "SERV".
    """

    # Common unit tokens seen in licitação/engenharia BOQs.
    STRONG_UNIT_TOKENS = {
        "UN", "UNID", "UND", "UNIDADE", "UNIDADES",
        "M", "M2", "M²", "M3", "M³", "CM", "MM", "KM",
        "L", "LT", "LITRO", "LITROS", "ML",
        "KG", "G", "TON", "T",
        "PAR", "PÇ", "PCA", "PECA", "PEÇA", "PEÇAS",
        "H", "HR", "HORA", "HORAS", "DIA", "DIAS", "MÊS", "MES",
        "SERV", "SERVICO", "SERVIÇO", "SERVIÇOS",
        "JOGO", "CONJUNTO", "KIT",
        "CX", "CAIXA", "PCT", "PACOTE", "FARDO", "ROLO", "EMB", "EMBALAGEM",
        "FRASCO", "GALAO", "GALÃO", "SACO", "BARRA",
    }
    QUALIFIER_WORDS = {"CX", "CAIXA", "PCT", "PACOTE", "KIT", "EMB", "EMBALAGEM", "FARDO", "ROLO", "JOGO", "CONJUNTO"}

    def _clean_unit(x: str) -> str:
        x = _norm(x)
        # Common in pdfplumber: "SERV\n." becomes "SERV ." after normalisation.
        x = x.replace(" .", ".")
        x = x.upper().strip()
        # Remove dots after abbreviations ("UN." -> "UN", "CX." -> "CX").
        x = re.sub(r"\b([A-Z]{1,4})\.", r"\1", x)
        # Normalise "C/ 50" -> "C/50".
        x = re.sub(r"\s*/\s*", "/", x)
        x = re.sub(r"\s+", " ", x).strip()
        return x

    def _looks_like_unit(cand: str) -> bool:
        if not cand:
            return False
        if "R$" in cand:
            return False
        if len(cand) > 40:
            return False
        if not re.search(r"[A-ZÀ-Ú]", cand):
            return False
        # Pure number (or decimal) isn't a unit.
        if re.fullmatch(r"\d+(?:[\.,]\d+)?", cand):
            return False
        # Reject obvious item-number strings.
        if _parse_item_id(cand) is not None:
            return False

        # If there are digits, accept common package qualifiers (e.g. "C/50",
        # "CAIXA 50 UNIDADE", "PACOTE 100 FOLHAS").
        if re.search(r"\d", cand):
            if re.search(r"\bC/\d+\b", cand):
                return True
            if any(w in cand for w in QUALIFIER_WORDS):
                return True
            # Some units include dimensions like "40 X 40 X 60 CM" (not a unit).
            return False

        return True

    obj_u = _clean_unit(obj_text) if obj_text else ""

    def _score_unit(cand: str, j: int) -> int:
        # Higher is better.
        score = 0

        # Strong signal: exact match to known unit tokens.
        if cand in STRONG_UNIT_TOKENS:
            score += 200

        # Short abbreviations are more likely units than long nouns.
        if re.fullmatch(r"[A-ZÀ-Ú]{1,4}[0-9]{0,2}", cand):
            score += 80

        # Package qualifiers ("CAIXA 50 UNIDADE", "C/50", etc).
        if re.search(r"\bC/\d+\b", cand):
            score += 70
        if any(w in cand for w in QUALIFIER_WORDS):
            score += 40

        # Prefer candidates near the expected unit column when we have one.
        if und_idx >= 0:
            score += max(0, 20 - abs(j - und_idx) * 5)

        # Avoid returning the object cell as unit when header mapping is missing.
        if obj_u and cand == obj_u and cand not in STRONG_UNIT_TOKENS:
            score -= 250

        # Long single-word nouns (e.g. "CONTAINER") are usually object names, not units.
        if " " not in cand and len(cand) >= 9 and cand not in STRONG_UNIT_TOKENS:
            score -= 60

        # Mild preference for shorter units.
        score -= min(len(cand), 30)

        return score

    def _best_from_indices(indices: list[int]) -> str:
        best: tuple[int, str] | None = None
        for j in indices:
            if 0 <= j < len(row):
                cand = _clean_unit(row[j])
                if not _looks_like_unit(cand):
                    continue
                sc = _score_unit(cand, j)
                if best is None or sc > best[0]:
                    best = (sc, cand)
        return best[1] if best is not None else ""

    # 1) Try mapped column directly.
    if 0 <= und_idx < len(row):
        cand = _clean_unit(row[und_idx])
        if _looks_like_unit(cand):
            return cand

    # 2) Window around expected column (or around start if und_idx=-1).
    if und_idx >= 0:
        window = list(range(max(0, und_idx - 3), min(len(row), und_idx + 4)))
    else:
        window = list(range(0, min(len(row), 6)))
    best = _best_from_indices(window)
    if best:
        return best

    # 3) Whole-row fallback (for tables without explicit headers).
    return _best_from_indices(list(range(len(row)))) or ""


def _estimate_confidence(item: ItemExtraido, doc_type: str) -> float:
    """Estimate how reliable an extracted item is on a 0.0–1.0 scale.

    Scoring:
      0.55 base
      +0.25 if quantity > 0
      +0.15 if unit is non-empty
      +0.05 if objeto is at least 20 characters

    The total is then multiplied by a doc_type weight:
      relacaoitens=1.0, edital=0.9, termo_referencia=0.8, others=0.85

    Args:
        item:     The partially constructed ItemExtraido.
        doc_type: Source document type key.

    Returns:
        Confidence score rounded to 2 decimal places, capped at 1.0.
    """
    base = 0.55
    if item.quantidade > 0:
        base += 0.25
    if item.unidade_fornecimento:
        base += 0.15
    if len(item.objeto) >= 20:
        base += 0.05
    w = {"relacaoitens": 1.0, "edital": 0.9, "termo_referencia": 0.8}.get(doc_type, 0.85)
    return round(min(1.0, base * w), 2)


# =============================================================================
# Metadata (heuristic) — shared by both parsers
# =============================================================================

# Two patterns covering the most common Pregão Eletrônico number formats:
#   Pattern 0: "Pregão Eletrônico Nº: 0001/2024"  (with optional separator/prefix)
#   Pattern 1: "PREGÃO ELETRÔNICO 000123"          (compact uppercase variant)
_PREGAO_PATTERNS = [
    re.compile(
        r"Preg[aã]o\s+Eletr[oô]nico\s*(n[ºo.]*)?\s*[:\-]?\s*"
        r"([0-9]{1,4}\s*/\s*[0-9]{2,4}|[0-9]{1,6})",
        re.I,
    ),
    re.compile(r"PREG[AÃ]O\s+ELETR[OÔ]NICO\s*([0-9]{1,6}(?:/[0-9]{2,4})?)", re.I),
]


def _extract_numero_pregao(text: str) -> str:
    """Extract the Pregão Eletrônico number from a block of text.

    Iterates over ``_PREGAO_PATTERNS`` and returns the first digit-containing
    capture group found (preferring the most specific group). Spaces within
    the number are removed to normalise formats like "0001 / 2024".

    Returns an empty string when no match is found.
    """
    for pat in _PREGAO_PATTERNS:
        m = pat.search(text)
        if m:
            # Reversed so that the most specific (last) group is tried first.
            for g in reversed(m.groups()):
                if g and re.search(r"\d", g):
                    return _norm(g).replace(" ", "")
            return _norm(m.group(0))
    return ""


def _extract_orgao_cidade_estado(text: str) -> tuple[str, str, str]:
    """Extract contracting authority, city and state from a block of text.

    Three heuristics, tried in order:
      1. "PREFEITURA MUNICIPAL DE <CITY>" — sets both orgao and cidade.
      2. "Município de <CITY>"            — fallback orgao/cidade.
      3. "<CITY>/<UF>" pattern            — extracts cidade and estado (UF).

    When cidade is still empty after all heuristics, it is inferred from the
    orgao string (everything after the last " de ").

    Returns:
        Tuple of (orgao, cidade, estado) — each is "" when not found.
    """
    orgao = cidade = estado = ""
    m = re.search(r"PREFEITURA\s+MUNICIPAL\s+DE\s+([A-ZÀ-Ú\s]+)", text, re.I)
    if m:
        city = _norm(m.group(1)).title()
        orgao = f"Prefeitura Municipal de {city}"
    m2 = re.search(r"Munic[ií]pio\s+de\s+([A-ZÀ-Ú\s]+)", text, re.I)
    if not orgao and m2:
        city = _norm(m2.group(1)).title()
        orgao = f"Município de {city}"
    # Matches "São Paulo/SP" — a common address notation in Brazilian documents.
    m3 = re.search(r"\b([A-ZÀ-Ú][A-Za-zÀ-ú\s'´`-]{2,})/([A-Z]{2})\b", text)
    if m3:
        cidade = _norm(m3.group(1)).title()
        estado = m3.group(2).upper()
    if not cidade and orgao:
        # Last resort: infer city from orgao, e.g. "Prefeitura Municipal de Goiânia" → "Goiânia"
        cidade = orgao.split(" de ", 1)[-1]
    return orgao, cidade, estado


# =============================================================================
# ── Relação de Itens text-block parser ───────────────────────────────────────
# =============================================================================

# Matches the item header line: "3 - Pasta eventos" or "10 – Serviço XYZ"
_RE_RI_ITEM_HEADER = re.compile(r"^(\d{1,5})\s*[-–]\s*(.+)$")

# Captures the "Descrição Detalhada" field value (greedy, DOTALL so it spans
# multiple lines — trimmed later by _parse_description_block).
_RE_RI_DESC      = re.compile(r"Descri[çc][aã]o\s+Detalhada\s*:\s*(.+)", re.I | re.DOTALL)

# Captures "Quantidade Total: 1.500" — dots and commas removed before int conversion.
_RE_RI_QTD_TOTAL = re.compile(r"Quantidade\s+Total\s*:\s*(\d[\d\.,]*)", re.I)

# Captures "Unidade de Fornecimento: Unidade" — stops at comma or newline.
_RE_RI_UNIDADE   = re.compile(r"Unidade\s+de\s+Fornecimento\s*:\s*([^\n\r,]+)", re.I)

# Captures lot labels: "LOTE 1", "LOTE 01", "LOTE ÚNICO", "LOTE UNICO"
_RE_RI_LOTE      = re.compile(r"\bLOTE\s+([A-Z]?\d{1,3}|ÚNICO|UNICO)\b", re.I)

# Marker used to detect whether a document is a Relação de Itens report.
_RI_MARKER = re.compile(r"RELA[ÇC][AÃ]O\s+DE\s+ITENS", re.I)


def _extract_text_from_pdf(file_path: Path) -> str:
    """Extract plain text from all pages of a PDF using pdfplumber.

    Pages are joined with a form-feed character (``\\f``) so that per-page
    boundaries are preserved for downstream heuristics if needed.

    Note: ``pdfplumber`` is imported inside this function because this code
    path was added after the module was originally written (which did not
    import pdfplumber at the top level — that was done by extractor.py).
    It should be moved to the top-level imports when convenient.

    Returns:
        Full document text as a single string.
    """
    import pdfplumber  # already a dependency — see module note above
    pages: list[str] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            pages.append(txt)
    return "\f".join(pages)


def _extract_text_from_docx(file_path: Path) -> str:
    """Extract plain text from a DOCX file using python-docx.

    Collects:
      - All paragraph texts in document order.
      - All table cell texts (pipe-separated per row) — needed for metadata
        detection and Relação de Itens auto-detection in DOCX attachments.

    Raises:
        RuntimeError: If python-docx is not installed.
    """
    try:
        from docx import Document  # python-docx
    except Exception as e:
        raise RuntimeError(
            "python-docx is required to parse .docx attachments. Install with `pip install python-docx`."
        ) from e

    doc = Document(str(file_path))
    parts: list[str] = []

    for p in doc.paragraphs:
        t = _norm(p.text)
        if t:
            parts.append(t)

    # Add table text too (helps metadata + relacaoitens detection)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [_norm(c.text) for c in row.cells]
            line = " | ".join([c for c in row_cells if c])
            if line:
                parts.append(line)

    return "\n".join(parts)


def _extract_tables_from_docx(file_path: Path, *, fonte_label: str) -> list[dict[str, Any]]:
    """Extract tables from a DOCX file using python-docx.

    Returns the same dict structure used by the PDF table extractor so that
    ``_extract_items_from_tables`` can process both PDF and DOCX tables
    through the same code path:
        [{'arquivo', 'pagina', 'indice_tabela', 'dados'}]

    ``pagina`` is set to ``None`` because DOCX tables have no page concept.

    Raises:
        RuntimeError: If python-docx is not installed.
    """
    try:
        from docx import Document  # python-docx
    except Exception as e:
        raise RuntimeError(
            "python-docx is required to parse .docx attachments. Install with `pip install python-docx`."
        ) from e

    doc = Document(str(file_path))
    out: list[dict[str, Any]] = []

    for ti, table in enumerate(doc.tables):
        dados: list[list[str]] = []
        for row in table.rows:
            cells = [_norm(cell.text) for cell in row.cells]
            if any(cells):
                dados.append(cells)

        if dados:
            out.append(
                {
                    "arquivo": fonte_label,
                    "pagina": None,       # DOCX tables have no page number
                    "indice_tabela": ti,
                    "dados": dados,
                }
            )

    return out


def _is_relacaoitens_text(text: str) -> bool:
    """Return True if the document text looks like a Relação de Itens report.

    Only inspects the first 2 000 characters to keep the check fast — the
    title is almost always in the document header.
    """
    return bool(_RI_MARKER.search(text[:2000]))


def _parse_description_block(block: str) -> str:
    """Extract the ``Descrição Detalhada`` value from a raw item block.

    The DOTALL regex captures everything from the label to the end of the block,
    but the description actually ends at the first line that begins with a
    recognised field keyword (e.g. "Tratamento Diferenciado:", "Quantidade:").
    This function trims the captured text at that boundary.

    Returns:
        Normalised description string, or "" if the field is absent.
    """
    m = _RE_RI_DESC.search(block)
    if not m:
        return ""
    raw = m.group(1)
    # Trim at the first sentinel line (e.g. "Tratamento Diferenciado: …")
    # We look for a newline followed by an uppercase keyword pattern.
    stop = re.search(
        r"\n(?:Tratamento|Aplicabilidade|Quantidade|Crit[eé]rio|Unidade|"
        r"Intervalo|Local|Valor\s+Unit)",
        raw,
        re.I,
    )
    if stop:
        raw = raw[: stop.start()]
    return _norm(raw)


def _extract_items_from_relacaoitens_text(
    text: str,
    *,
    doc_type: str,
    fonte: Optional[str],
    debug: bool,
) -> tuple[list[ItemExtraido], dict[str, str]]:
    """Parse a full Relação de Itens document (all pages joined) into items.

    Strategy:
      1. Detect the document-level lote from the first ~500 characters.
      2. Split all lines into per-item blocks using ``_RE_RI_ITEM_HEADER`` as
         the delimiter. A new block starts only when the item number is
         plausibly sequential (avoids false positives from date stamps like
         "14/08/2024" which would match the header regex).
      3. For each block, extract description, quantity, unit and lote via regex.
      4. Deduplicate by item number, keeping the highest-scoring version.
      5. Extract document metadata (pregão number, orgao, cidade, estado).

    Args:
        text:     Full document text from ``_extract_text_from_pdf`` or
                  ``_extract_text_from_docx``.
        doc_type: Used for confidence estimation weighting.
        fonte:    Source file label; stored on each item when ``debug=True``.
        debug:    Controls whether ``fonte`` is populated.

    Returns:
        Tuple of (items list, metadata dict).
    """
    itens: list[ItemExtraido] = []
    current_lote: Optional[str] = None

    # Detect lote in the first ~500 chars (document header area)
    m_lote = _RE_RI_LOTE.search(text[:500])
    if m_lote:
        val = m_lote.group(1).upper()
        current_lote = _format_lote(val)

    lines = text.splitlines()

    # Split lines into blocks, each starting at an item header line.
    # We collect (item_num, item_name, block_lines) triples.
    blocks: list[tuple[int, str, list[str]]] = []
    current_num: Optional[int] = None
    current_name: str = ""
    current_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        m = _RE_RI_ITEM_HEADER.match(stripped)
        if m:
            num = int(m.group(1))
            name = _norm(m.group(2))
            # Only start a new block if the number is plausibly sequential.
            # This avoids grabbing date stamps like "14/08/2024" → num=14.
            if not blocks or num == blocks[-1][0] + 1 or (num > 0 and num <= 9999):
                if current_num is not None:
                    blocks.append((current_num, current_name, current_lines))
                current_num = num
                current_name = name
                current_lines = []
                continue
        if current_num is not None:
            current_lines.append(line)

    if current_num is not None:
        blocks.append((current_num, current_name, current_lines))

    for num, name, blines in blocks:
        block_text = "\n".join(blines)

        # Description: prefer the 'Descrição Detalhada' field; fall back to name
        descricao = _parse_description_block(block_text)
        objeto = descricao if descricao else name

        # Quantity: remove thousand separators before converting to int
        m_qtd = _RE_RI_QTD_TOTAL.search(block_text)
        quantidade = int(m_qtd.group(1).replace(".", "").replace(",", "")) if m_qtd else 0

        # Unit: title-cased for consistent formatting
        m_und = _RE_RI_UNIDADE.search(block_text)
        unidade = _norm(m_und.group(1)).title() if m_und else ""

        # Per-block lote override — rare in single-lote documents but needed
        # for multi-lote Relação de Itens where each block declares its lot.
        m_lote2 = _RE_RI_LOTE.search(block_text[:200])
        lote = current_lote
        if m_lote2:
            val2 = m_lote2.group(1).upper()
            lote = _format_lote(val2)

        it = ItemExtraido(
            item=str(num),
            objeto=objeto,
            quantidade=quantidade,
            unidade_fornecimento=unidade,
            lote=lote,
            fonte=fonte if debug else None,
        )
        it.confianca = _estimate_confidence(it, doc_type)
        itens.append(it)

    # Deduplicate by item number (same item may appear on multiple pages).
    # Keep the version with the highest (confidence, has_qty, objeto_length) score.
    best: dict[str, ItemExtraido] = {}
    for it in itens:
        cur = best.get(it.item)
        if cur is None:
            best[it.item] = it
        else:
            def _score(x: ItemExtraido) -> tuple:
                return (x.confianca, 1 if x.quantidade else 0, len(x.objeto))
            if _score(it) > _score(cur):
                best[it.item] = it

    final = sorted(best.values(), key=lambda x: x.item_sort_key())

    # Extract document-level metadata for the merge_metadata step.
    numero = _extract_numero_pregao(text)
    orgao, cidade, estado = _extract_orgao_cidade_estado(text)
    meta = {"numero_pregao": numero, "orgao": orgao, "cidade": cidade, "estado": estado}

    return final, meta


# =============================================================================
# Table-based parser
# =============================================================================

def _table_looks_like_continuation(rows: list[list[str]]) -> bool:
    """Return True if the table appears to continue items from a previous table.

    IMPORTANT: some continuation pages start with one or more rows that *don't*
    contain an item number (e.g. a wrapped description line like
    "FORNECIMENTO DE MATERIAL."). For that reason we must scan a small window
    of rows, not just the first non-empty row.

    Heuristic: if the first non-empty row's first 8 cells contain at least one
    parseable item identifier, the table is likely a continuation page.
    This allows the parser to reuse the header map from the previous table
    rather than requiring every page to repeat the column headers.
    """

    for row in rows[:5]:
        if not any(row):
            continue
        for c in row[:8]:
            if _parse_item_id(c) is not None:
                return True
    return False

def _merge_header_rows(header_rows: list[list[str]]) -> list[str]:
    """Merge multi-line header blocks into a single synthetic header row.

    Some PDFs draw the table header across several stacked rows (e.g. the word
    'ITEM' split as 'ITE' + 'M', and 'DESCRIÇÃO' placed on a different header row).
    pdfplumber then yields several rows with many empty/None cells.

    This helper concatenates non-empty cells column-wise to produce one row that
    `_detect_header_indices` can understand.
    """
    if not header_rows:
        return []
    ncols = max(len(r) for r in header_rows)
    merged: list[str] = []
    for j in range(ncols):
        parts: list[str] = []
        for r in header_rows:
            if j < len(r):
                v = _norm(r[j])
                if v:
                    parts.append(v)
        merged.append(" ".join(parts))
    return merged


def _extract_items_from_tables(
    tables_json: list[dict[str, Any]],
    *,
    doc_type: str,
    debug: bool,
) -> tuple[list[ItemExtraido], dict[str, str]]:
    """Extract items from a list of table dicts (PDF or DOCX origin).

    This version fixes two common lote-related failure modes:

      1) **Lote headers mid-table**:
         Some PDFs concatenate multiple lots into a single extracted table, e.g.
         the table starts with the tail of the previous lot, then a row like
         "LOTE 09 PRINCIPAL", then a new header + items.
         Lote detection must therefore be **row-based and sequential**, not
         a single "first rows" guess for the whole table.

      2) **Continuation rows before a repeated header**:
         A table can begin with data rows (continuation) and only later repeat
         the header row. We initialise `active_header_map` from the best header
         we can find (merged header block or a later header row) so we can
         parse those early rows too, while still skipping the header row when
         we encounter it.

    Processing pipeline per table (sequential scan):
      - Update `current_lote` whenever a lote-title row is found.
      - Update `active_header_map` whenever a header row is found.
      - Parse item rows using `active_header_map`.
      - Append non-item rows to the previous item's `objeto` when they look
        like wrapped description lines.

    Returns:
        Tuple of (items list, metadata dict).
    """
    itens: list[ItemExtraido] = []
    anexos: set[str] = set()

    current_lote: Optional[str] = None
    last_item: Optional[ItemExtraido] = None
    last_header_map: Optional[dict[str, int]] = None

    for t in tables_json:
        fonte_arquivo = _norm(t.get("arquivo", ""))
        if fonte_arquivo:
            anexos.add(fonte_arquivo)

        rows = _flatten_table(t.get("dados", []))
        if not rows:
            continue

        # ── Pick an initial header_map to parse *all* rows in this table ──
        header_map: Optional[dict[str, int]] = None

        # (1) Multi-row header: detect header block above the first data row.
        first_data_idx: Optional[int] = None
        for ri, row in enumerate(rows[:80]):
            if any(_parse_item_id(c) is not None for c in row[:8]):
                first_data_idx = ri
                break

        if first_data_idx is not None and first_data_idx > 0:
            merged_header = _merge_header_rows(rows[max(0, first_data_idx - 10): first_data_idx])
            hm = _detect_header_indices(merged_header)
            if "item" in hm and "obj" in hm:
                header_map = hm
                last_header_map = hm

        # (2) Fallback: find any header row within the first 60 rows.
        if header_map is None:
            for row in rows[:60]:
                hm = _row_is_header(row)
                if hm:
                    header_map = hm
                    last_header_map = hm
                    break

        # (3) Continuation: if still none, reuse previous header_map if table starts with items.
        if header_map is None and last_header_map is not None:
            if _table_looks_like_continuation(rows):
                header_map = last_header_map

        if header_map is None:
            continue  # no structure we can parse

        active_header_map = header_map

        # ── Sequential scan: update lote/header mid-table and parse rows ──
        for row in rows:
            if not any(row):
                continue

            # (A) Lote-title row (can appear at the top OR mid-table).
            lote_label = _extract_lote_from_row(row)
            if lote_label:
                current_lote = lote_label
                last_item = None
                continue

            # (B) Header row (can also appear mid-table).
            hm = _row_is_header(row)
            if hm:
                active_header_map = hm
                last_header_map = hm
                last_item = None
                continue

            # (C) Parse data row.
            item_id: Optional[tuple[int, Optional[str]]] = None
            for c in row[:8]:
                parsed = _parse_item_id(c)
                if parsed is not None:
                    item_id = parsed
                    break

            if item_id is not None:
                major, sub = item_id
                item_str = f"{major}.{sub}" if sub is not None else str(major)

                obj_idx = active_header_map.get("obj", 0)
                qtd_idx = active_header_map.get("qtd", -1)
                und_idx = active_header_map.get("und", -1)

                qty_cell = row[qtd_idx] if 0 <= qtd_idx < len(row) else ""
                und_cell = row[und_idx] if 0 <= und_idx < len(row) else ""
                obj_text = _pick_obj(row, obj_idx)

                # Skip "section rows" that look like group/lot dividers.
                if not qty_cell and not und_cell and _is_section_title(obj_text):
                    last_item = None
                    continue

                obj = obj_text.strip()
                qtd = _extract_qty_with_header(row, qtd_idx, major)
                und = _extract_unit_with_header(row, und_idx, obj).strip()

                it = ItemExtraido(
                    item=item_str,
                    objeto=obj,
                    quantidade=qtd,
                    unidade_fornecimento=und,
                    lote=current_lote,
                    fonte=(fonte_arquivo or None) if debug else None,
                )
                it.confianca = _estimate_confidence(it, doc_type)
                itens.append(it)
                last_item = it
            else:
                # Non-item row: may be a wrapped description line for the previous item.
                if not last_item:
                    continue
                txt = _pick_obj(row, active_header_map.get("obj", 0)).strip()
                if not txt or len(txt) <= 2:
                    continue
                if _is_section_title(txt):
                    last_item = None
                    continue
                # Avoid appending obvious header-ish noise if something slipped through.
                if _row_is_header(row):
                    last_item = None
                    continue
                if _extract_lote_from_row(row):
                    last_item = None
                    continue
                last_item.objeto = (last_item.objeto + " " + txt).strip()

    # Deduplicate by (lote, item) — keep the highest-scoring version.
    best: dict[tuple, ItemExtraido] = {}
    for it in itens:
        key = (it.lote, it.item)
        cur = best.get(key)
        if cur is None:
            best[key] = it
        else:
            def _score(x: ItemExtraido) -> tuple:
                return (x.confianca, 1 if x.quantidade else 0, len(x.objeto))
            if _score(it) > _score(cur):
                best[key] = it

    final_itens = [v for v in best.values() if v.objeto or v.quantidade]
    final_itens.sort(key=lambda x: (x.lote or "", x.item_sort_key()))

    # Metadata from all table text.
    all_text = "\n".join([_table_text(_flatten_table(t.get("dados", []))) for t in tables_json])
    numero = _extract_numero_pregao(all_text)
    orgao, cidade, estado = _extract_orgao_cidade_estado(all_text)
    meta = {"numero_pregao": numero, "orgao": orgao, "cidade": cidade, "estado": estado}
    return final_itens, meta


# =============================================================================
# Public dataclass & entrypoint
# =============================================================================

@dataclass
class ParsedAttachment:
    """Result of parsing a single attachment file (PDF or DOCX).

    Attributes:
        doc_type:          Document type key (e.g. "edital", "relacaoitens").
        items:             Extracted and deduplicated list of items.
        meta:              Metadata dict with keys numero_pregao/orgao/cidade/estado.
        tables_json_path:  Path to the intermediate tables JSON (None for text path).
        parsed_json_path:  Path to the final resultado JSON written for this attachment.
    """
    doc_type: str
    items: list[ItemExtraido]
    meta: dict[str, str]
    tables_json_path: Optional[Path] = None
    parsed_json_path: Optional[Path] = None


def parse_attachment(
    file_path: Path,
    doc_type: str,
    *,
    project_root: Path,
    tables_out_dir: Path,
    parsed_out_dir: Path,
    debug: bool = False,
) -> ParsedAttachment:
    """Full extraction pipeline for one PDF or DOCX attachment.

    Routing logic:
      1. If ``doc_type == "relacaoitens"`` → text-block parser.
      2. If auto-detection finds "RELAÇÃO DE ITENS" in the first 2 000 chars
         of the extracted text → text-block parser.
      3. Otherwise → table-based parser (original behaviour).

    Both paths write intermediate and final JSON files under the provided
    output directories and return a ``ParsedAttachment`` with the results.

    Args:
        file_path:      Path to the attachment to parse (.pdf or .docx).
        doc_type:       Hint about the document type (may trigger text parser).
        project_root:   Used for computing relative paths in outputs.
        tables_out_dir: Where to write ``*_tables.json`` (table path only).
        parsed_out_dir: Where to write ``*_resultado.json``.
        debug:          When True, ``fonte`` fields are populated on items.

    Returns:
        ``ParsedAttachment`` with items, metadata and output paths.

    Raises:
        ValueError: For unsupported file extensions.
    """
    tables_out_dir.mkdir(parents=True, exist_ok=True)
    parsed_out_dir.mkdir(parents=True, exist_ok=True)

    fonte_label = _safe_relpath(file_path, project_root)
    suffix = file_path.suffix.lower()

    # ── Route: text-block parser for Relação de Itens ──────────────────────
    use_text_parser = doc_type == "relacaoitens"

    if not use_text_parser and suffix in (".pdf", ".docx"):
        # Auto-detect: peek at page text before committing to a parser.
        # Exceptions are silently ignored — we fall through to the table parser.
        try:
            peek = _extract_text_from_pdf(file_path) if suffix == ".pdf" else _extract_text_from_docx(file_path)
            if _is_relacaoitens_text(peek):
                use_text_parser = True
        except Exception:
            pass  # fall through to table parser

    if use_text_parser and suffix in (".pdf", ".docx"):
        full_text = _extract_text_from_pdf(file_path) if suffix == ".pdf" else _extract_text_from_docx(file_path)
        items, meta = _extract_items_from_relacaoitens_text(
            full_text,
            doc_type=doc_type,
            fonte=fonte_label,
            debug=debug,
        )

        parsed_json_path = parsed_out_dir / f"{file_path.stem}_resultado.json"
        res = ResultadoLicitacao(
            arquivo_json=parsed_json_path.name,
            numero_pregao=meta.get("numero_pregao", ""),
            orgao=meta.get("orgao", ""),
            cidade=meta.get("cidade", ""),
            estado=meta.get("estado", ""),
            anexos_processados=[fonte_label],
            itens_extraidos=items,
        )
        parsed_json_path.write_text(res.to_json(debug=debug), encoding="utf-8")

        return ParsedAttachment(
            doc_type=doc_type,
            items=items,
            meta=meta,
            tables_json_path=None,       # no tables JSON for the text-block path
            parsed_json_path=parsed_json_path,
        )

    # ── Route: table-based parser (PDF via extractor; DOCX via python-docx) ─
    if suffix == ".pdf":
        tables = process_file(file_path, project_root=project_root)
    elif suffix == ".docx":
        tables = _extract_tables_from_docx(file_path, fonte_label=fonte_label)
    else:
        raise ValueError(f"Unsupported attachment type: {suffix}")

    tables_json_path = tables_out_dir / f"{file_path.stem}_tables.json"
    write_tables_json(tables, tables_json_path)

    items, meta = _extract_items_from_tables(tables, doc_type=doc_type, debug=debug)

    parsed_json_path = parsed_out_dir / f"{file_path.stem}_resultado.json"
    res = ResultadoLicitacao(
        arquivo_json=tables_json_path.name,
        numero_pregao=meta.get("numero_pregao", ""),
        orgao=meta.get("orgao", ""),
        cidade=meta.get("cidade", ""),
        estado=meta.get("estado", ""),
        anexos_processados=[fonte_label],
        itens_extraidos=items,
    )
    parsed_json_path.write_text(res.to_json(debug=debug), encoding="utf-8")

    return ParsedAttachment(
        doc_type=doc_type,
        items=items,
        meta=meta,
        tables_json_path=tables_json_path,
        parsed_json_path=parsed_json_path,
    )


# =============================================================================
# Batch CLI (table path only — same as before)
# =============================================================================

def process_tables_json_file(
    path: Path, *, doc_type: str = "edital", debug: bool = False
) -> ResultadoLicitacao:
    """Process a pre-extracted ``*_tables.json`` file through the table parser.

    Useful for re-running the item extraction step without re-extracting tables
    from the original PDF, e.g. when tuning heuristics.

    Args:
        path:     Path to a ``*_tables.json`` file.
        doc_type: Document type hint for confidence weighting.
        debug:    Controls whether ``fonte`` fields are included.

    Returns:
        A ``ResultadoLicitacao`` with extracted items and metadata.
    """
    tables = json.loads(path.read_text(encoding="utf-8"))
    items, meta = _extract_items_from_tables(tables, doc_type=doc_type, debug=debug)
    return ResultadoLicitacao(
        arquivo_json=path.name,
        numero_pregao=meta.get("numero_pregao", ""),
        orgao=meta.get("orgao", ""),
        cidade=meta.get("cidade", ""),
        estado=meta.get("estado", ""),
        anexos_processados=list(
            {t.get("arquivo", "") for t in tables if t.get("arquivo")}
        ),
        itens_extraidos=items,
    )


def cli_main() -> None:
    """CLI entry point for batch-processing pre-extracted table JSON files.

    Reads all files matching ``--pattern`` under ``--input-dir``, runs them
    through the table parser, writes per-file result JSONs, and produces a
    consolidated output file.

    Example:
        python -m parsers.pdf_parser \\
            --input-dir outputs/tabelas \\
            --output-dir outputs/parsed \\
            --doc-type edital
    """
    ap = argparse.ArgumentParser()
    ap.add_argument("--input-dir", type=Path, required=True)
    ap.add_argument("--output-dir", type=Path, required=True)
    ap.add_argument("--pattern", type=str, default="*_tables.json")
    ap.add_argument("--doc-type", type=str, default="edital")
    ap.add_argument("--debug", action="store_true")
    args = ap.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    results: list[dict] = []
    for p in sorted(args.input_dir.glob(args.pattern)):
        res = process_tables_json_file(p, doc_type=args.doc_type, debug=args.debug)
        results.append(res.to_dict(debug=args.debug))
        out_file = args.output_dir / (p.stem.replace("_tables", "") + "_resultado.json")
        out_file.write_text(res.to_json(debug=args.debug), encoding="utf-8")

    consolidated = args.output_dir / "resultado_licitacoes_consolidado.json"
    consolidated.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    cli_main()